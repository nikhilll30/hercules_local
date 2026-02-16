import os
import shutil  # Added to delete old memory folders
import pandas as pd
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# --- LANGCHAIN IMPORTS (The Switch to Ollama) ---
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.tools import Tool
from langchain_core.tools.retriever import create_retriever_tool

# --- TOOLS ---
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_experimental.tools import PythonAstREPLTool

# --- LANGGRAPH ---
from langgraph.prebuilt import create_react_agent
from langgraph.checkpoint.memory import MemorySaver

# --- CONFIGURATION ---
import streamlit as st

# Check if the key exists in the secrets file
if "TAVILY_API_KEY" in st.secrets:
    os.environ["TAVILY_API_KEY"] = st.secrets["TAVILY_API_KEY"]



# --- ENGINE SWAP: GOOGLE -> OLLAMA ---
print("--- CONNECTING TO LOCAL OLLAMA ---")

# 1. The Brain (Llama 3.1)
# base_url is usually localhost:11434
llm = ChatOllama(model="llama3.1", temperature=0)

# 2. The Embeddings (Nomic)
embeddings = OllamaEmbeddings(model="nomic-embed-text")

memory = MemorySaver()


# --- FILE GENERATION TOOLS (Unchanged) ---
def generate_word_doc(content: str, filename: str = "report.docx") -> str:
    try:
        doc = DocxDocument()
        doc.add_heading('Hercules Local Report', 0)
        doc.add_paragraph(content)
        path = f"./{filename}" if not filename.endswith(".docx") else filename
        doc.save(path)
        return f"SUCCESS: Word doc saved at {path}"
    except Exception as e:
        return f"Error: {e}"


def generate_pdf_doc(content: str, filename: str = "report.pdf") -> str:
    try:
        path = f"./{filename}" if not filename.endswith(".pdf") else filename
        c = canvas.Canvas(path, pagesize=letter)
        y = 750
        for line in content.split('\n'):
            if y < 40: c.showPage(); y = 750
            c.drawString(40, y, line)
            y -= 14
        c.save()
        return f"SUCCESS: PDF saved at {path}"
    except Exception as e:
        return f"Error: {e}"


def generate_excel_sheet(data_csv_string: str, filename: str = "data.xlsx") -> str:
    import io
    try:
        df = pd.read_csv(io.StringIO(data_csv_string))
        path = f"./{filename}" if not filename.endswith(".xlsx") else filename
        df.to_excel(path, index=False)
        return f"SUCCESS: Excel saved at {path}"
    except Exception as e:
        return f"Error: {e}"


tools_export = [
    Tool(name="create_word_doc", func=generate_word_doc, description="Create Word file."),
    Tool(name="create_pdf_doc", func=generate_pdf_doc, description="Create PDF file."),
    Tool(name="create_excel_file", func=generate_excel_sheet, description="Create Excel from CSV.")
]


# --- AGENT BUILDER ---
def get_agent_with_docs(file_paths):
    print("--- HERCULES LOCAL INITIALIZING ---")

    text_docs = []
    csv_dataframes = {}

    for file_path in file_paths:
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".pdf":
                text_docs.extend(PyPDFLoader(file_path).load())
            elif ext == ".txt":
                text_docs.extend(TextLoader(file_path).load())
            elif ext == ".csv":
                name = os.path.basename(file_path).replace(".", "_")
                df = pd.read_csv(file_path)
                csv_dataframes[name] = df
        except Exception as e:
            print(f"Error loading {file_path}: {e}")

    tools = []

    # 1. VECTOR STORE (Local Nomic Embeddings)
    if text_docs:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(text_docs)

        # --- MEMORY WIPE FIX ---
        # This deletes the old database folder so the agent forgets previous files.
        db_path = "./hercules_db_ollama"

        # Check if the folder exists and delete it
        if os.path.exists(db_path):
            try:
                shutil.rmtree(db_path)
                print(f"--- CLEARED OLD MEMORY AT {db_path} ---")
            except Exception as e:
                print(f"Warning: Could not delete old memory: {e}")

        # Create new vector store in the clean folder
        vector_store = Chroma.from_documents(
            documents=splits,
            embedding=embeddings,
            collection_name="hercules_local",
            persist_directory=db_path
        )

        retriever_tool = create_retriever_tool(
            vector_store.as_retriever(),
            "search_documents",
            "Search uploaded files."
        )
        tools.append(retriever_tool)

    # 2. PYTHON TOOL
    if csv_dataframes:
        python_tool = PythonAstREPLTool(locals={"dfs": csv_dataframes})
        python_tool.name = "python_data_analyst"
        python_tool.description = "Analyze CSVs in 'dfs'. Use matplotlib for charts (save as plot.png)."
        tools.append(python_tool)

    # 3. WEB SEARCH (Still needs internet)
    try:
        search_tool = TavilySearchResults(max_results=3)
        tools.append(search_tool)
    except:
        pass

    tools.extend(tools_export)

    system_instruction = (
        "You are Hercules Local, running on Ollama. "
        "1. Use 'search_documents' for PDFs. "
        "2. Use 'python_data_analyst' for CSVs. "
        "3. Use 'tavily_search_results_json' for web info. "
        "4. Be concise."
    )

    agent = create_react_agent(llm, tools, checkpointer=memory)
    return agent, system_instruction
